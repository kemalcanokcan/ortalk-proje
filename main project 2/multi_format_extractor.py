import os
import json
import csv
import re
import logging
from datetime import datetime
from pdf_extractor import PDFExtractor

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


class MultiFormatExtractor:
    """
    A multi-format invoice extractor that can handle PDF, DOCX, JSON, CSV, and Excel files
    """
    
    def __init__(self, file_path):
        self.file_path = file_path
        self.file_extension = os.path.splitext(file_path)[1].lower()
        # Common aliases to help header/field mapping across formats
        self.field_aliases = {
            'invoice_number': ['invoice_number', 'invoiceNumber', 'invoice_no', 'number', 'no', 'fatura_no', 'fatura no', 'invoice id', 'id'],
            'invoice_date': ['invoice_date', 'invoiceDate', 'date', 'issue_date', 'tarih', 'fatura tarihi'],
            'vendor_name': ['vendor_name', 'seller_name', 'supplier_name', 'from_name', 'satici_adi', 'firma', 'company', 'accounting supplier'],
            'vendor_tax_id': ['vendor_tax_id', 'seller_tax_id', 'supplier_vkn', 'from_vkn', 'satici_vkn', 'vkn', 'tax id', 'taxid', 'vergino'],
            'customer_name': ['customer_name', 'buyer_name', 'to_name', 'alici_adi', 'customer', 'musteri', 'accounting customer'],
            'customer_tax_id': ['customer_tax_id', 'buyer_tax_id', 'to_vkn', 'alici_vkn', 'vkn', 'tax id', 'vergino'],
            'total_amount': ['total_amount', 'total', 'grand_total', 'toplam', 'genel toplam'],
            'subtotal': ['subtotal', 'sub_total', 'net_total', 'ara_toplam'],
            'tax_amount': ['tax_amount', 'vat_amount', 'kdv_tutari', 'kdv tutarı', 'kdv'],
            'currency': ['currency', 'para_birimi', 'para birimi', 'birim', 'curr', 'currency code']
        }
        
    def extract_invoice_data(self):
        """
        Extract invoice data based on file format
        
        Returns:
            dict: Structured invoice data
        """
        logger.info(f"Extracting data from {self.file_extension} file: {self.file_path}")
        
        try:
            if self.file_extension == '.pdf':
                return self._extract_from_pdf()
            elif self.file_extension == '.docx':
                return self._extract_from_docx()
            elif self.file_extension == '.json':
                return self._extract_from_json()
            elif self.file_extension == '.csv':
                return self._extract_from_csv()
            elif self.file_extension in ('.xlsx', '.xls'):
                return self._extract_from_excel()
            else:
                raise ValueError(f"Unsupported file format: {self.file_extension}")
                
        except Exception as e:
            logger.error(f"Error extracting data from {self.file_extension} file: {e}")
            raise
    
    def _extract_from_pdf(self):
        """Extract data from PDF. Use PDFExtractor, fallback to OCR if needed."""
        try:
            extractor = PDFExtractor(self.file_path)
            data = extractor.extract_invoice_data()
        except Exception as e:
            logger.warning(f"Primary PDF extraction failed: {e}. Trying OCR fallback.")
            data = {}

        # Check minimal completeness; if weak, try OCR
        needs_ocr = False
        if not data or (
            not data.get('invoice_number') and
            not data.get('line_items') and
            not data.get('total_amount')
        ):
            needs_ocr = True

        if needs_ocr:
            ocr_data = self._extract_from_pdf_with_ocr()
            # Merge: prefer structured fields when present
            data = {**ocr_data, **{k: v for k, v in (data or {}).items() if v}}

        return data

    def _extract_from_pdf_with_ocr(self):
        """OCR fallback using pdf2image + pytesseract, then parse text."""
        try:
            from pdf2image import convert_from_path
            import pytesseract
            from PIL import Image
        except ImportError:
            logger.error("OCR dependencies missing. Install: pip install pdf2image pytesseract pillow")
            return {}

        try:
            # Convert pages to images
            images = convert_from_path(self.file_path, dpi=300)
            text_pages = []
            for img in images:
                # Basic pre-processing can improve OCR
                ocr_text = pytesseract.image_to_string(img, lang='tur+eng')
                text_pages.append(ocr_text)
            text_content = "\n\n".join(text_pages)

            # Parse OCRed text with existing text parser
            parsed = self._parse_text_content(text_content, tables_content=None)
            logger.info("OCR-based PDF extraction completed")
            return parsed
        except Exception as e:
            logger.error(f"OCR extraction failed: {e}")
            return {}
    
    def _extract_from_docx(self):
        """Extract data from DOCX file"""
        try:
            from docx import Document
            
            # Read the document
            doc = Document(self.file_path)
            text_content = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extract text from tables
            tables_content = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                tables_content.append(table_data)
            
            # Use similar parsing logic as PDF extractor
            return self._parse_text_content(text_content, tables_content)
            
        except ImportError:
            logger.error("python-docx library not available. Install it with: pip install python-docx")
            raise
        except Exception as e:
            logger.error(f"Error extracting from DOCX: {e}")
            raise
    
    def _extract_from_json(self):
        """Extract data from JSON file"""
        try:
            with open(self.file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Normalize JSON structure to our standard format
            invoice_data = {
                'invoice_number': '',
                'invoice_date': '',
                'vendor_name': '',
                'vendor_tax_id': '',
                'vendor_address': '',
                'customer_name': '',
                'customer_tax_id': '',
                'customer_address': '',
                'line_items': [],
                'subtotal': '',
                'tax_amount': '',
                'total_amount': '',
                'notes': ''
            }
            
            # Map common JSON field names to our structure
            field_mappings = {
                'invoice_number': ['invoice_number', 'invoiceNumber', 'invoice_no', 'number', 'no', 'fatura_no'],
                'invoice_date': ['invoice_date', 'invoiceDate', 'date', 'issue_date', 'tarih'],
                'vendor_name': ['vendor_name', 'seller_name', 'supplier_name', 'from_name', 'satici_adi'],
                'vendor_tax_id': ['vendor_tax_id', 'seller_tax_id', 'supplier_vkn', 'from_vkn', 'satici_vkn'],
                'vendor_address': ['vendor_address', 'seller_address', 'supplier_address', 'from_address', 'satici_adres'],
                'customer_name': ['customer_name', 'buyer_name', 'to_name', 'alici_adi'],
                'customer_tax_id': ['customer_tax_id', 'buyer_tax_id', 'to_vkn', 'alici_vkn'],
                'customer_address': ['customer_address', 'buyer_address', 'to_address', 'alici_adres'],
                'total_amount': ['total_amount', 'total', 'grand_total', 'toplam'],
                'subtotal': ['subtotal', 'sub_total', 'net_total', 'ara_toplam'],
                'tax_amount': ['tax_amount', 'vat_amount', 'kdv_tutari'],
                'notes': ['notes', 'comments', 'notlar']
            }
            
            # Extract main fields
            for our_field, possible_fields in field_mappings.items():
                for field in possible_fields:
                    if field in data and data[field]:
                        invoice_data[our_field] = str(data[field]).strip()
                        break
            
            # Extract line items
            line_items_keys = ['line_items', 'items', 'products', 'kalemler', 'urunler']
            for key in line_items_keys:
                if key in data and isinstance(data[key], list):
                    for item in data[key]:
                        if isinstance(item, dict):
                            line_item = self._normalize_line_item(item)
                            if line_item:
                                invoice_data['line_items'].append(line_item)
                    break
            
            # Normalize date
            if invoice_data['invoice_date']:
                invoice_data['invoice_date'] = self._normalize_date(invoice_data['invoice_date'])
            # Clean numeric totals
            for k in ('subtotal', 'tax_amount', 'total_amount'):
                if invoice_data.get(k):
                    invoice_data[k] = self._clean_number(invoice_data[k])
            
            # Clean addresses
            if invoice_data.get('vendor_address'):
                invoice_data['vendor_address'] = self._clean_address(invoice_data['vendor_address'])
            if invoice_data.get('customer_address'):
                invoice_data['customer_address'] = self._clean_address(invoice_data['customer_address'])
            
            logger.info(f"Extracted data from JSON with {len(invoice_data['line_items'])} line items")
            return invoice_data
            
        except Exception as e:
            logger.error(f"Error extracting from JSON: {e}")
            raise
    
    def _extract_from_csv(self):
        """Extract data from CSV file"""
        try:
            invoice_data = {
                'invoice_number': '',
                'invoice_date': '',
                'vendor_name': '',
                'vendor_tax_id': '',
                'vendor_address': '',
                'customer_name': '',
                'customer_tax_id': '',
                'customer_address': '',
                'line_items': [],
                'subtotal': '',
                'tax_amount': '',
                'total_amount': '',
                'notes': ''
            }
            
            with open(self.file_path, 'r', encoding='utf-8') as f:
                # Try to detect CSV format
                sample = f.read(1024)
                f.seek(0)
                
                # Use csv.Sniffer to detect delimiter
                sniffer = csv.Sniffer()
                delimiter = sniffer.sniff(sample).delimiter
                
                reader = csv.DictReader(f, delimiter=delimiter)
                rows = list(reader)
            
            if not rows:
                raise ValueError("CSV file is empty or has no data rows")
            
            # Check if this is a line items CSV or a header+items CSV
            first_row = rows[0]
            
            # Look for invoice header information in the first few rows
            header_fields = ['invoice_number', 'invoice_date', 'vendor_name', 'customer_name']
            has_header_info = any(field in first_row for field in header_fields)
            
            if has_header_info:
                # Extract header information from first row or metadata
                for key, value in first_row.items():
                    key_lower = key.lower().replace(' ', '_').replace('-', '_')
                    if 'invoice' in key_lower and ('number' in key_lower or 'no' in key_lower):
                        invoice_data['invoice_number'] = str(value).strip()
                    elif 'date' in key_lower or 'tarih' in key_lower:
                        invoice_data['invoice_date'] = str(value).strip()  # Keep exact format
                    elif 'vendor' in key_lower or 'seller' in key_lower or 'satici' in key_lower:
                        if 'name' in key_lower or 'adi' in key_lower:
                            invoice_data['vendor_name'] = str(value).strip()
                        elif 'vkn' in key_lower or 'tax' in key_lower:
                            invoice_data['vendor_tax_id'] = str(value).strip()
                    elif 'customer' in key_lower or 'buyer' in key_lower or 'alici' in key_lower:
                        if 'name' in key_lower or 'adi' in key_lower:
                            invoice_data['customer_name'] = str(value).strip()
                        elif 'vkn' in key_lower or 'tax' in key_lower:
                            invoice_data['customer_tax_id'] = str(value).strip()
            
            # Extract line items from all rows
            for row in rows:
                line_item = self._normalize_csv_line_item(row)
                if line_item:
                    invoice_data['line_items'].append(line_item)
            
            # Calculate totals if missing
            if not invoice_data['total_amount'] and invoice_data['line_items']:
                total = 0.0
                for item in invoice_data['line_items']:
                    amt = item.get('amount', '')
                    total += float(self._clean_number(amt)) if amt else 0.0
                invoice_data['total_amount'] = str(total)

            # Normalize date if present
            if invoice_data.get('invoice_date'):
                invoice_data['invoice_date'] = self._normalize_date(invoice_data['invoice_date'])

            # Clean numeric totals if present
            for k in ('subtotal', 'tax_amount', 'total_amount'):
                if invoice_data.get(k):
                    invoice_data[k] = self._clean_number(invoice_data[k])
            
            logger.info(f"Extracted data from CSV with {len(invoice_data['line_items'])} line items")
            return invoice_data
            
        except Exception as e:
            logger.error(f"Error extracting from CSV: {e}")
            raise
    
    def _extract_from_excel(self):
        """Extract data from Excel (.xlsx/.xls) using pandas."""
        try:
            import pandas as pd
        except ImportError:
            logger.error("pandas/openpyxl not available. Install: pip install pandas openpyxl")
            raise

        invoice_data = {
            'invoice_number': '',
            'invoice_date': '',
            'vendor_name': '',
            'vendor_tax_id': '',
            'vendor_address': '',
            'customer_name': '',
            'customer_tax_id': '',
            'customer_address': '',
            'line_items': [],
            'subtotal': '',
            'tax_amount': '',
            'total_amount': '',
            'notes': ''
        }

        try:
            # Try to read all sheets
            xls = pd.ExcelFile(self.file_path)
            # Heuristic: header info in first sheet, items in sheet named like 'items' or the one with many rows
            items_sheet_name = None
            for name in xls.sheet_names:
                lname = name.lower()
                if any(k in lname for k in ['item', 'lines', 'kalem', 'ürün']):
                    items_sheet_name = name
                    break

            # Parse header from first sheet
            df0 = xls.parse(xls.sheet_names[0])
            # Coerce columns to lowercase normalized
            df0_cols = {c: str(c).strip().lower() for c in df0.columns}
            for our, aliases in self.field_aliases.items():
                for alias in aliases:
                    for col, norm in df0_cols.items():
                        if alias == norm:
                            val = df0[col].iloc[0] if not df0.empty else ''
                            if pd.notna(val):
                                invoice_data[our] = str(val).strip()
                                break
                    if invoice_data.get(our):
                        break

            # Items
            if items_sheet_name:
                df_items = xls.parse(items_sheet_name)
            else:
                # Fallback: if first sheet has many rows and typical item columns
                df_items = df0.copy()

            # Normalize item columns
            norm_cols = {i: str(c).strip().lower() for i, c in enumerate(df_items.columns)}
            # Identify columns
            idx_desc = idx_qty = idx_price = idx_vat = idx_amount = -1
        
            for i, name in norm_cols.items():
                if any(word in name for word in ['description', 'item', 'product', 'açıklama', 'ürün']):
                    idx_desc = i
                elif any(word in name for word in ['quantity', 'qty', 'miktar', 'adet']):
                    idx_qty = i
                elif any(word in name for word in ['price', 'unit', 'birim', 'fiyat']):
                    idx_price = i
                elif any(word in name for word in ['vat', 'tax', 'kdv', 'vergi']):
                    idx_vat = i
                elif any(word in name for word in ['total', 'amount', 'toplam', 'tutar']):
                    idx_amount = i

            for _, row in df_items.iterrows():
                item = {}
                if idx_desc is not None:
                    v = row.iloc[idx_desc]
                    if pd.notna(v) and str(v).strip():
                        item['description'] = str(v).strip()
                if idx_qty is not None:
                    v = row.iloc[idx_qty]
                    if pd.notna(v) and str(v).strip():
                        item['quantity'] = str(v).strip()
                if idx_price is not None:
                    v = row.iloc[idx_price]
                    if pd.notna(v) and str(v).strip():
                        item['unit_price'] = str(v).strip()
                if idx_vat is not None:
                    v = row.iloc[idx_vat]
                    if pd.notna(v) and str(v).strip():
                        item['tax_rate'] = str(v).strip()
                if idx_amount is not None:
                    v = row.iloc[idx_amount]
                    if pd.notna(v) and str(v).strip():
                        item['amount'] = str(v).strip()
                if item.get('description'):
                    if 'quantity' not in item:
                        item['quantity'] = '1'
                    invoice_data['line_items'].append(item)

            # Normalize date and numeric totals
            if invoice_data.get('invoice_date'):
                invoice_data['invoice_date'] = self._normalize_date(str(invoice_data['invoice_date']))
            for k in ('subtotal', 'tax_amount', 'total_amount'):
                if invoice_data.get(k):
                    invoice_data[k] = self._clean_number(invoice_data[k])

            logger.info(f"Extracted data from Excel with {len(invoice_data['line_items'])} line items")
            return invoice_data
        except Exception as e:
            logger.error(f"Error extracting from Excel: {e}")
            raise
    
    def _parse_text_content(self, text_content, tables_content=None):
        """Parse text content using similar logic to PDF extractor"""
        invoice_data = {
            'invoice_number': '',
            'invoice_date': '',
            'vendor_name': '',
            'vendor_tax_id': '',
            'vendor_address': '',
            'customer_name': '',
            'customer_tax_id': '',
            'customer_address': '',
            'line_items': [],
            'subtotal': '',
            'tax_amount': '',
            'total_amount': '',
            'notes': ''
        }
        
        # Extract invoice number
        invoice_patterns = [
            r'(?:Invoice\s+Number|Invoice\s+No|Fatura\s+No|No)[:\s]+([A-Za-z0-9\-._/]+)',
            r'(?:Number|No)[:\s]+([A-Za-z0-9\-._/]+)'
        ]
        
        for pattern in invoice_patterns:
            match = re.search(pattern, text_content, re.IGNORECASE)
            if match:
                invoice_data['invoice_number'] = match.group(1).strip()
                break
        
        # Extract date
        date_patterns = [
            r'(?:Invoice\s+Date|Date|Tarih)[:\s]+(\d{1,2}[./\-]\d{1,2}[./\-]\d{2,4})',
            r'(\d{1,2}[./\-]\d{1,2}[./\-]\d{2,4})'
        ]
        
        for pattern in date_patterns:
            match = re.search(pattern, text_content, re.IGNORECASE)
            if match:
                invoice_data['invoice_date'] = self._normalize_date(match.group(1).strip())
                break
        
        # Extract vendor/customer names (simplified)
        company_pattern = r'([A-Z][A-Za-z\s]*(?:Ltd|Inc|Corp|AŞ|LTD|ŞTİ)[A-Za-z\s]*)'
        companies = re.findall(company_pattern, text_content)
        
        if len(companies) >= 2:
            invoice_data['vendor_name'] = companies[0].strip()
            invoice_data['customer_name'] = companies[1].strip()
        elif len(companies) == 1:
            invoice_data['vendor_name'] = companies[0].strip()
        
        # Extract totals exactly as written
        total_patterns = [
            r'(?:Total|Grand\s+Total|Toplam)[:\s]+([0-9.,₺$€\s]+)',
            r'([0-9.,₺$€\s]+)\s*(?:TL|₺|$|€)'
        ]
        
        for pattern in total_patterns:
            matches = re.findall(pattern, text_content, re.IGNORECASE)
            if matches:
                # Take the first match and keep it exactly as found
                for match in matches:
                    if match.strip():
                        invoice_data['total_amount'] = match.strip()
                        break
                break
        
        # Extract line items from tables if available
        if tables_content:
            for table in tables_content:
                items = self._extract_items_from_table(table)
                invoice_data['line_items'].extend(items)
        
        return invoice_data
    
    def _extract_items_from_table(self, table_data):
        """Extract line items from table data"""
        items = []
        
        if not table_data or len(table_data) < 2:
            return items
        
        # Assume first row is header
        headers = [h.lower().strip() for h in table_data[0]]
        
        # Map header indices
        desc_idx = quantity_idx = price_idx = vat_idx = total_idx = -1
        
        for i, header in enumerate(headers):
            if any(word in header for word in ['description', 'item', 'product', 'açıklama', 'ürün']):
                desc_idx = i
            elif any(word in header for word in ['quantity', 'qty', 'miktar', 'adet']):
                quantity_idx = i
            elif any(word in header for word in ['price', 'unit', 'birim', 'fiyat']):
                price_idx = i
            elif any(word in header for word in ['vat', 'tax', 'kdv', 'vergi']):
                vat_idx = i
            elif any(word in header for word in ['total', 'amount', 'toplam', 'tutar']):
                total_idx = i
        
        # Extract items from data rows
        for row in table_data[1:]:
            if len(row) <= max(desc_idx, quantity_idx, price_idx, vat_idx, total_idx):
                continue
            
            item = {}
            
            if desc_idx >= 0 and desc_idx < len(row):
                item['description'] = row[desc_idx].strip()
            
            if quantity_idx >= 0 and quantity_idx < len(row):
                item['quantity'] = row[quantity_idx].strip()
            else:
                item['quantity'] = '1'
            
            if price_idx >= 0 and price_idx < len(row):
                item['unit_price'] = row[price_idx].strip()
            
            if vat_idx >= 0 and vat_idx < len(row):
                item['tax_rate'] = row[vat_idx].strip()  # Keep exact format
            
            if total_idx >= 0 and total_idx < len(row):
                item['amount'] = row[total_idx].strip()
            
            # Validate item
            if item.get('description') and len(item['description']) > 2:
                items.append(item)
        
        return items
    
    def _normalize_line_item(self, item_dict):
        """Normalize a line item from JSON format"""
        line_item = {}
        
        # Map common field names
        field_mappings = {
            'description': ['description', 'item', 'product', 'name', 'açıklama', 'ürün'],
            'quantity': ['quantity', 'qty', 'miktar', 'adet'],
            'unit_price': ['unit_price', 'price', 'rate', 'birim_fiyat', 'fiyat'],
            'tax_rate': ['tax_rate', 'vat_rate', 'vat', 'kdv', 'kdv_orani'],
            'amount': ['amount', 'total', 'toplam', 'tutar']
        }
        
        for our_field, possible_fields in field_mappings.items():
            for field in possible_fields:
                if field in item_dict and item_dict[field] is not None:
                    # Keep all values exactly as they appear
                    line_item[our_field] = str(item_dict[field]).strip()
                    break
        
        # Only set quantity default if missing
        if 'quantity' not in line_item:
            line_item['quantity'] = '1'
        # Don't set tax_rate default - leave empty if not found
        
        return line_item if line_item.get('description') else None
    
    def _normalize_csv_line_item(self, row_dict):
        """Normalize a line item from CSV row"""
        line_item = {}
        
        for key, value in row_dict.items():
            if not value or not value.strip():
                continue
            
            key_lower = key.lower().replace(' ', '_').replace('-', '_')
            
            if any(word in key_lower for word in ['description', 'item', 'product', 'name', 'açıklama', 'ürün']):
                line_item['description'] = value.strip()
            elif any(word in key_lower for word in ['quantity', 'qty', 'miktar', 'adet']):
                line_item['quantity'] = value.strip()  # Keep exact format
            elif any(word in key_lower for word in ['price', 'unit', 'birim', 'fiyat']):
                line_item['unit_price'] = value.strip()  # Keep exact format
            elif any(word in key_lower for word in ['vat', 'tax', 'kdv', 'vergi']):
                line_item['tax_rate'] = value.strip()  # Keep exact format
            elif any(word in key_lower for word in ['total', 'amount', 'toplam', 'tutar']):
                line_item['amount'] = value.strip()  # Keep exact format
        
        # Only set quantity default if missing
        if 'quantity' not in line_item:
            line_item['quantity'] = '1'
        # Don't set tax_rate default - leave empty if not found
        
        return line_item if line_item.get('description') else None
    
    def _normalize_date(self, date_str):
        """Normalize date to ISO format (YYYY-MM-DD)"""
        if not date_str:
            return datetime.now().strftime('%Y-%m-%d')
        
        # Check if already in ISO format
        if re.match(r'^\d{4}-\d{2}-\d{2}$', date_str):
            return date_str
        
        # Try to parse different date formats
        date_patterns = [
            r'(\d{1,2})[./\-](\d{1,2})[./\-](\d{4})',  # DD/MM/YYYY
            r'(\d{1,2})[./\-](\d{1,2})[./\-](\d{2})',   # DD/MM/YY
            r'(\d{4})[./\-](\d{1,2})[./\-](\d{1,2})'    # YYYY/MM/DD
        ]
        
        for pattern in date_patterns:
            match = re.search(pattern, date_str)
            if match:
                parts = match.groups()
                try:
                    if len(parts[0]) == 4:  # YYYY/MM/DD format
                        year, month, day = parts
                    else:  # DD/MM/YYYY or DD/MM/YY format
                        day, month, year = parts
                        if len(year) == 2:
                            year = '20' + year if int(year) < 50 else '19' + year
                    
                    return f"{year}-{month.zfill(2)}-{day.zfill(2)}"
                except:
                    continue
        
        logger.warning(f"Could not parse date '{date_str}', using current date")
        return datetime.now().strftime('%Y-%m-%d')
    
    def _parse_vat_rate(self, vat_str, default_rate='18'):
        """Parse VAT rate - same logic as in other classes"""
        # Import the logic from pdf_extractor or duplicate it here
        # For now, simplified version
        if not vat_str:
            return default_rate
        
        vat_str = str(vat_str).strip().replace('%', '')
        
        try:
            vat_float = float(vat_str)
            if 0 <= vat_float <= 30:
                return str(vat_float)
            elif vat_float > 100:
                # Likely a parsing error, try to correct
                if 100 <= vat_float < 1000:
                    return str(vat_float / 10)
                elif vat_float >= 1000:
                    return str(vat_float / 100)
            return default_rate
        except:
            return default_rate
    
    def _clean_number(self, number_str):
        """Clean and standardize number format"""
        if not number_str:
            return '0'
        
        # Remove currency symbols and other non-numeric characters
        cleaned = re.sub(r'[^\d.,]', '', str(number_str))
        
        # Handle Turkish number format
        if ',' in cleaned and '.' in cleaned:
            # 1.234,56 format
            cleaned = cleaned.replace(',', '.')
            parts = cleaned.split('.')
            if len(parts) > 2:
                cleaned = ''.join(parts[:-1]) + '.' + parts[-1]
        elif ',' in cleaned:
            # Replace comma with dot for decimal
            cleaned = cleaned.replace(',', '.')
        
        try:
            return str(float(cleaned))
        except:
            return '0'
    
    def _clean_address(self, address_str):
        """Clean address string by removing prefixes and fixing spacing."""
        if not address_str or not isinstance(address_str, str):
            return ""

        # 1. Remove common prefixes, case-insensitive
        prefixes = [
            r'SATICI ADRESİ\s*:', r'VENDOR ADDRESS\s*:', r'ADRES\s*:', r'ADDRESS\s*:',
            r'VERGİ DAİRESİ\s*:', r'MERSİS NO\s*:', r'VD\s*:'
        ]
        for prefix in prefixes:
            address_str = re.sub(f'^{prefix}', '', address_str, flags=re.IGNORECASE).strip()

        # 2. Insert space between letter and digit (e.g., "Mahallesi3028" -> "Mahallesi 3028")
        address_str = re.sub(r'([a-zA-ZıİğĞüÜşŞöÖçÇ])(\d)', r'\1 \2', address_str)
        # 3. Insert space between digit and letter (e.g., "CADDE16A" -> "CADDE 16A")
        address_str = re.sub(r'(\d)([a-zA-ZıİğĞüÜşŞöÖçÇ])', r'\1 \2', address_str)

        # 4. Normalize whitespace
        address_str = ' '.join(address_str.split())

        return address_str
